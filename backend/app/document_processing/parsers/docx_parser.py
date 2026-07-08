"""
DealOS AI — DOCX Parser.

Extracts text from Word documents using python-docx.
"""

import structlog
from docx import Document as DocxDocument

logger = structlog.get_logger(__name__)


class DocxParser:
    """Extract text content from Word documents."""

    @staticmethod
    async def parse(content: bytes) -> list[dict]:
        """
        Parse a DOCX file and extract text.

        Returns a single-page representation since DOCX files
        don't have inherent page boundaries.
        """
        import io
        pages = []

        try:
            doc = DocxDocument(io.BytesIO(content))
            paragraphs = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)

            if paragraphs:
                full_text = "\n\n".join(paragraphs)
                # Split into pseudo-pages of ~3000 chars for manageable chunks
                page_size = 3000
                for i in range(0, len(full_text), page_size):
                    page_text = full_text[i:i + page_size]
                    pages.append({
                        "page_number": (i // page_size) + 1,
                        "text": page_text,
                        "metadata": {"source_type": "docx"},
                    })

            logger.info("docx_parsed", paragraphs=len(paragraphs), pages=len(pages))

        except Exception as e:
            logger.error("docx_parse_error", error=str(e))
            raise

        return pages
